#!/usr/bin/env python3
"""Generate the final StudyFlow Assignment 3 Word report."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "StudyFlow_Assignment3_Report.docx"
GRAFANA = ROOT / "monitoring" / "screenshots" / "grafana-dashboard.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D9DEE5"
TEXT = "263238"
MUTED = "5F6B76"
GREEN = "177245"
WHITE = "FFFFFF"
CONTENT_TWIPS = 9360  # 6.5 inches with 1-inch Letter margins


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, twips=CONTENT_TWIPS):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(twips))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_bottom_border(paragraph, color=BLUE, size=14, space=5):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_hyperlink(paragraph, text, url, color=BLUE, underline=True):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    run_props.append(run_color)
    if underline:
        run_underline = OxmlElement("w:u")
        run_underline.set(qn("w:val"), "single")
        run_props.append(run_underline)
    run.append(run_props)
    run_text = OxmlElement("w:t")
    run_text.text = text
    run.append(run_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(NAVY)
    code_style.paragraph_format.left_indent = Inches(0.16)
    code_style.paragraph_format.right_indent = Inches(0.16)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "StudyFlow Assignment 3 - Performance, Security, and Monitoring"
    doc.core_properties.subject = "CSCI 4177/5709 Assignment 3"
    doc.core_properties.author = "Arup Halder"
    doc.core_properties.keywords = "StudyFlow, JMeter, OWASP ZAP, Prometheus, Grafana"
    doc.core_properties.comments = "Final Brightspace submission report"

    configure_headers_and_footers(section)


def configure_headers_and_footers(section):
    # The editorial cover carries the document identity, so subsequent pages
    # use intentionally quiet page furniture: no running header and a simple
    # page number. Explicitly populate all variants for consistent rendering.
    for header in (section.first_page_header, section.header, section.even_page_header):
        header.paragraphs[0].text = ""

    for footer in (section.first_page_footer, section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(p)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_section_title(doc, number, title, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(kicker.upper())
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    heading = doc.add_heading(f"{number}. {title}", level=1)
    if not kicker:
        heading.paragraph_format.page_break_before = True


def add_bullets(doc, items, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        p.paragraph_format.space_after = Pt(2 if compact else 4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.27)
        p.paragraph_format.first_line_indent = Inches(-0.17)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_callout(doc, heading, body, color=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_grid(table, [CONTENT_TWIPS])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, CONTENT_TWIPS)
    set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
    set_cell_shading(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(heading)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.runs[0].font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    p.add_run(text)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GREY)
    p_pr.append(shd)
    return p


def add_table(doc, headers, rows, widths, font_size=9, header_fill=LIGHT_GREY, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    twips = [round(w * 1440) for w in widths]
    set_table_grid(table, twips)

    header = table.rows[0]
    set_repeat_table_header(header)
    keep_row_together(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_width(cell, twips[index])
        set_cell_margins(cell)
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        if aligns:
            p.alignment = aligns[index]
        run = p.add_run(value)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_width(cell, twips[index])
            set_cell_margins(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if aligns:
                p.alignment = aligns[index]
            run = p.add_run(str(value))
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(28)
    add_bottom_border(p, color=BLUE, size=28, space=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("CSCI 4177/5709  |  ASSIGNMENT 3")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("StudyFlow")
    run.font.name = "Calibri"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("Performance, Security, and Monitoring Report")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_grid(table, [CONTENT_TWIPS])
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_width(cell, CONTENT_TWIPS)
    set_cell_margins(cell, top=150, start=120, bottom=150, end=120)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("FINAL SUBMISSION")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p = cell.add_paragraph(
        "Verified against the assignment instructions and rubric. Includes reproducible "
        "load tests, four measured optimizations, security remediation, and live monitoring evidence."
    )
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Prepared by")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Arup Halder")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Course and term")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph("CSCI 4177/5709, Summer 2026")
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Repository")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    add_hyperlink(p, "github.com/beaprogram/studyflow-assignment3", "https://github.com/beaprogram/studyflow-assignment3")

    p = doc.add_paragraph("Report date: July 15, 2026")
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def add_executive_summary(doc):
    add_section_title(doc, "1", "Executive Summary", "Outcome")
    doc.add_paragraph(
        "StudyFlow was profiled under light and moderate load, optimized on both the client "
        "and server, rescanned for security weaknesses, and instrumented for continuous "
        "operational visibility. All required implementation and evidence artifacts are in "
        "the linked repository."
    )
    add_callout(
        doc,
        "Measured outcome",
        "Overall average latency improved by 42.3% under light load and 47.0% under "
        "moderate load. p95 latency improved by 46.2% and 51.3%, respectively, while all "
        "4,808 submitted JMeter samples completed with a 0.00% error rate.",
    )
    doc.add_heading("Rubric coverage", level=2)
    add_table(
        doc,
        ["Rubric area", "Completed evidence", "Status"],
        [
            ["Baseline testing", "JMX plus light/moderate baseline CSVs and bottleneck analysis", "Complete"],
            ["Client optimization", "Route code splitting and memoized course rendering", "Complete"],
            ["Server optimization", "TTL response cache and compound MongoDB index", "Complete"],
            ["Performance comparison", "Same plan rerun; avg, p95, RPS, errors, and percentages", "Complete"],
            ["Security", "ZAP before/after reports and two remediations", "Complete"],
            ["Monitoring", "Prometheus metrics and six-panel Grafana dashboard", "Complete"],
            ["Documentation", "This report, repository instructions, and evidence inventory", "Complete"],
        ],
        [1.35, 4.30, 0.85],
        font_size=8.4,
    )
    doc.add_heading("Submission at a glance", level=2)
    add_bullets(
        doc,
        [
            "Repository: https://github.com/beaprogram/studyflow-assignment3",
            "Primary Brightspace file: StudyFlow_Assignment3_Report.pdf",
            "Editable backup: StudyFlow_Assignment3_Report.docx",
        ],
        compact=True,
    )


def add_methodology(doc):
    add_section_title(doc, "2", "Baseline Test Method", "JMeter")
    doc.add_paragraph(
        "The baseline and optimized measurements use the same parameterized JMeter sampler "
        "tree. A setup thread logs in once and extracts the bearer token. The load group then "
        "requests the authenticated course list, index page, JavaScript entry bundle, and CSS "
        "bundle with a 500 ms Constant Timer. This preserves an identical workload before and after."
    )
    doc.add_heading("Load scenarios", level=2)
    add_table(
        doc,
        ["Scenario", "Virtual users", "Ramp-up", "Loops/user", "Samples/run"],
        [
            ["Light", "10", "30 seconds", "10", "401"],
            ["Moderate", "50", "60 seconds", "10", "2,001"],
        ],
        [1.30, 1.15, 1.25, 1.30, 1.50],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    doc.add_heading("Metrics captured", level=2)
    add_bullets(
        doc,
        [
            "Average response time and p95 latency for each sampler and for the total run.",
            "Throughput in requests per second and request error rate.",
            "Raw timestamp, elapsed time, status, bytes, latency, connect time, URL, and thread counts in CSV.",
        ],
        compact=True,
    )
    doc.add_heading("Three baseline bottlenecks", level=2)
    add_table(
        doc,
        ["Bottleneck", "Baseline evidence", "Why it mattered"],
        [
            ["POST Login", "324 ms light; 296 ms moderate", "Slowest single request; includes authentication and database work."],
            ["GET Courses", "77.6/78.6 ms avg; 86/87 ms p95", "Repeated authenticated query dominated latency during the load group."],
            ["Main JS bundle", "578,126 bytes", "Large initial payload contained charting code not needed on first render."],
        ],
        [1.35, 2.10, 3.05],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Interpretation note",
        "Login runs once during setup, so it was identified as a bottleneck but was not the "
        "focus of the repeated read-path optimization. The course endpoint and first-load "
        "bundle offered the highest impact within the assignment scope.",
        color=LIGHT_GREY,
    )


def add_client_optimizations(doc):
    add_section_title(doc, "3", "Client-Side Optimizations", "React and Vite")
    doc.add_heading("3.1 Route-level code splitting", level=2)
    doc.add_paragraph(
        "The Analytics page imports Recharts, the heaviest client dependency. It was changed "
        "from an eager import to React.lazy with a dynamic import and Suspense fallback. The "
        "browser now downloads chart code only when the user navigates to Analytics."
    )
    add_code(doc, "const Analytics = lazy(() => import('./pages/Analytics.jsx'))")
    add_table(
        doc,
        ["Bundle metric", "Baseline", "Optimized", "Improvement"],
        [["Initial JavaScript entry", "578,126 bytes", "177,350 bytes", "69.3% smaller"]],
        [2.05, 1.45, 1.45, 1.55],
    )
    doc.add_paragraph(
        "The optimized build emits the Analytics/Recharts code as a separate deferred chunk. "
        "This reduces initial network transfer, parsing, and evaluation work for users who do "
        "not immediately open the analytics route."
    )

    doc.add_heading("3.2 Memoized course list and derived data", level=2)
    doc.add_paragraph(
        "CourseRow now uses React.memo, summary calculations and the row collection use "
        "useMemo, and the loading callback uses useCallback. Stable course data no longer "
        "causes every row to render after unrelated parent-state changes."
    )
    add_code(doc, "const CourseRow = memo(function CourseRow({ course }) { ... })")
    add_table(
        doc,
        ["Deterministic benchmark", "Baseline", "Optimized", "Improvement"],
        [["20 rows + 50 unrelated updates", "1,020 renders", "20 renders", "98.0% fewer"]],
        [2.55, 1.25, 1.25, 1.45],
    )
    doc.add_paragraph(
        "The benchmark is included at client/benchmarks/memoization-benchmark.mjs and runs "
        "with npm run benchmark:memo. It isolates render behavior, so the result is reproducible "
        "without relying on device-specific browser timing."
    )
    add_callout(
        doc,
        "Client impact",
        "The first optimization reduces bytes delivered on initial navigation; the second "
        "reduces repeated rendering after the page is loaded. Together they address network, "
        "JavaScript evaluation, and React reconciliation costs.",
    )


def add_server_optimizations(doc):
    add_section_title(doc, "4", "Server-Side Optimizations", "Express and MongoDB")
    doc.add_heading("4.1 Short-lived course-list cache", level=2)
    doc.add_paragraph(
        "Read-heavy course-list payloads are cached in process memory for 30 seconds. Keys "
        "include the authenticated user, page, limit, and term. Course create, update, and delete "
        "operations immediately invalidate that user's list entries, limiting stale data risk."
    )
    add_code(doc, "const key = `courses:list:${userId}:${page}:${limit}:${term}`;\ncache.set(key, payload, 30 * 1000);")
    add_table(
        doc,
        ["GET Courses", "Baseline avg", "Optimized avg", "Improvement", "Errors"],
        [
            ["Light", "77.6 ms", "40.0 ms", "48.5%", "0.00%"],
            ["Moderate", "78.6 ms", "38.3 ms", "51.3%", "0.00%"],
        ],
        [1.15, 1.30, 1.35, 1.35, 1.35],
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    doc.add_paragraph(
        "An X-Cache response header exposes HIT, MISS, or DISABLED, making the behavior easy to "
        "verify. CACHE_ENABLED=false preserves a controlled baseline path."
    )

    doc.add_heading("4.2 Compound query-and-sort index", level=2)
    doc.add_paragraph(
        "The course-list query filters by owner and sorts by createdAt descending. The model now "
        "defines { owner: 1, createdAt: -1 }, allowing MongoDB to return records in index order "
        "instead of scanning and sorting the collection."
    )
    add_table(
        doc,
        ["Explain metric", "Collection scan", "Compound index", "Improvement"],
        [
            ["Plan", "SORT <- COLLSCAN", "LIMIT <- FETCH <- IXSCAN", "No in-memory sort"],
            ["Documents examined", "10,006", "20", "99.8% fewer"],
            ["Execution time", "11 ms", "1 ms", "90.9% faster"],
        ],
        [1.55, 1.55, 1.85, 1.55],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Why both changes are needed",
        "The index improves unavoidable database reads and cold-cache requests. The cache removes "
        "repeated database work during short bursts. Write-triggered invalidation keeps both "
        "performance and correctness in balance.",
    )


def add_results(doc):
    add_section_title(doc, "5", "Before-and-After Comparison", "Measured results")
    doc.add_paragraph(
        "The optimized build was rerun with the same JMX plan, user counts, ramp times, loops, "
        "sampler tree, and think time. Percentage improvement for latency is calculated as "
        "(baseline - optimized) / baseline x 100."
    )
    doc.add_heading("Overall results", level=2)
    add_table(
        doc,
        ["Scenario / metric", "Baseline", "Optimized", "Change"],
        [
            ["Light average latency", "22.6 ms", "13.0 ms", "42.3% better"],
            ["Light p95 latency", "80 ms", "43 ms", "46.2% better"],
            ["Light throughput", "8.25 req/s", "8.32 req/s", "+0.9%"],
            ["Light error rate", "0.00%", "0.00%", "No errors"],
            ["Moderate average latency", "21.7 ms", "11.5 ms", "47.0% better"],
            ["Moderate p95 latency", "78 ms", "38 ms", "51.3% better"],
            ["Moderate throughput", "24.92 req/s", "25.03 req/s", "+0.4%"],
            ["Moderate error rate", "0.00%", "0.00%", "No errors"],
        ],
        [2.40, 1.35, 1.35, 1.40],
        font_size=8.7,
    )
    doc.add_heading("Endpoint-level results", level=2)
    add_table(
        doc,
        ["Scenario / endpoint", "Avg before", "Avg after", "p95 before", "p95 after", "Avg gain"],
        [
            ["Light - GET Courses", "77.6", "40.0", "86", "47", "48.5%"],
            ["Light - JS bundle", "4.2", "3.3", "6", "5", "21.9%"],
            ["Moderate - GET Courses", "78.6", "38.3", "87", "44", "51.3%"],
            ["Moderate - JS bundle", "3.6", "2.7", "5", "4", "25.9%"],
        ],
        [1.85, 0.95, 0.95, 0.90, 0.90, 0.95],
        font_size=8.1,
        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("Units: ")
    r.bold = True
    p.add_run("All endpoint average and p95 values are milliseconds.")
    add_callout(
        doc,
        "Result interpretation",
        "Latency improved substantially while throughput rose slightly. Because each scenario "
        "uses fixed loops and a 500 ms timer, throughput is primarily workload-bound rather than "
        "an open-ended saturation measure. The unchanged 0.00% error rate confirms stability.",
        color=LIGHT_GREY,
    )


def add_security(doc):
    add_section_title(doc, "6", "OWASP ZAP Security Validation", "Security")
    doc.add_paragraph(
        "A headless ZAP baseline scan was captured before remediation and repeated against the "
        "optimized local production preview after the fixes. Both HTML and JSON reports are "
        "stored under zap/."
    )
    doc.add_paragraph(
        "The headless ZAP baseline scan reported 9 alerts: 0 high, 2 medium, 5 low, and "
        "2 informational. Both medium-severity findings are remediated below with "
        "configuration snippets. The five low-severity alerts (Cross-Origin-Embedder-Policy, "
        "Cross-Origin-Opener-Policy, Cross-Origin-Resource-Policy, Permissions-Policy, and "
        "X-Content-Type-Options headers) are addressed by the same security-header changes "
        "as defense in depth. The informational alerts require no remediation."
    )
    doc.add_heading("Baseline findings and fixes", level=2)
    add_table(
        doc,
        ["Baseline alert", "Risk / count", "Remediation", "Final result"],
        [
            ["CSP Header Not Set", "Medium / 3", "Restrictive Content-Security-Policy across Vite, Netlify, and serve", "Resolved"],
            ["Missing Anti-clickjacking Header", "Medium / 1", "frame-ancestors 'none' plus X-Frame-Options: DENY", "Resolved"],
        ],
        [1.55, 1.05, 2.75, 1.15],
        font_size=8.4,
    )
    doc.add_heading("Remediation detail", level=2)
    add_code(
        doc,
        "default-src 'self'; script-src 'self'; style-src 'self'; img-src 'self' data:;\n"
        "connect-src 'self'; object-src 'none'; base-uri 'self'; form-action 'self';\n"
        "frame-ancestors 'none'",
    )
    add_bullets(
        doc,
        [
            "Removed the only inline React style, so style-src does not require 'unsafe-inline'.",
            "Added X-Frame-Options: DENY and frame-ancestors 'none' for legacy and modern clickjacking protection.",
            "Added nosniff, no-referrer, Permissions-Policy, and cross-origin isolation headers as defense in depth.",
        ],
        compact=True,
    )
    doc.add_heading("Final scan", level=2)
    add_callout(
        doc,
        "0 high-risk and 0 medium-risk findings",
        "The July 15, 2026 after-scan crawled six URLs. Its only remaining alerts are informational: "
        "Modern Web Application and Storable but Non-Cacheable Content. The two required medium "
        "findings are absent from the final JSON and HTML reports. The content alert changed from "
        "'Storable and Cacheable' to 'Storable but Non-Cacheable' because the new Cache-Control "
        "headers took effect.",
    )
    doc.add_paragraph(
        "Evidence files: zap/zap-report-before-web.html, zap/zap-report-before-web.json, "
        "zap/zap-report-after-web.html, zap/zap-report-after-web.json, and zap/remediation.md."
    )


def add_monitoring(doc):
    add_section_title(doc, "7", "Prometheus and Grafana Monitoring", "Observability")
    doc.add_paragraph(
        "The API exposes Prometheus metrics at /metrics using prom-client. Docker Compose runs "
        "Prometheus and Grafana; provisioning files connect the data source and load the StudyFlow "
        "dashboard automatically. The dashboard was verified with live JMeter traffic."
    )
    add_table(
        doc,
        ["Required signal", "Implementation / panel"],
        [
            ["CPU", "Node process user and system CPU rate"],
            ["Memory", "Resident memory and V8 heap used"],
            ["Response time", "p50, p95, and p99 from the HTTP duration histogram"],
            ["Error rate", "4xx/5xx share derived from request counters"],
            ["Additional", "Throughput by route and requests in flight"],
        ],
        [1.65, 4.85],
        font_size=8.6,
    )
    if not GRAFANA.exists():
        raise FileNotFoundError(f"Missing monitoring screenshot: {GRAFANA}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    picture = p.add_run().add_picture(str(GRAFANA), width=Inches(6.18))
    picture._inline.docPr.set(
        "descr",
        "Grafana dashboard showing StudyFlow API CPU, memory, latency, error rate, throughput, and requests in flight.",
    )
    p = doc.add_paragraph("Figure 1. StudyFlow Grafana dashboard under live API load.", style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_validation(doc):
    add_section_title(doc, "8", "Validation and Reproducibility", "Quality assurance")
    doc.add_heading("Final validation status", level=2)
    add_table(
        doc,
        ["Check", "Result", "Evidence"],
        [
            ["API smoke tests", "6 of 6 passing", "api/tests/smoke.test.js"],
            ["Client production build", "Passing", "Vite optimized build"],
            ["Production dependency audit", "0 vulnerabilities", "npm audit --omit=dev"],
            ["Memoization benchmark", "98.0% fewer row renders", "client/benchmarks/"],
            ["JMeter runs", "4,808 samples; 0.00% errors", "Four CSV result files"],
            ["MongoDB explain", "IXSCAN; 20 docs vs 10,006", "api/seed/explain-courses.js"],
            ["OWASP ZAP", "0 high; 0 medium", "Final JSON/HTML reports"],
            ["Monitoring", "Prometheus up; six panels", "Dashboard JSON and screenshot"],
        ],
        [2.10, 1.85, 2.55],
        font_size=8.3,
    )
    doc.add_heading("Core reproduction commands", level=2)
    add_code(
        doc,
        "cd api && npm test\n"
        "cd client && npm run build && npm run benchmark:memo\n"
        "cd jmeter && python3 compare.py\n"
        "cd monitoring && docker compose up -d",
    )
    doc.add_heading("Repository deliverable inventory", level=2)
    add_bullets(
        doc,
        [
            "client/ - optimized React/Vite source, benchmark, and security-header configuration.",
            "api/ - optimized Express/MongoDB source, tests, index evidence, cache, and metrics.",
            "jmeter/ - local and deployment JMX plans plus baseline/optimized light/moderate CSVs.",
            "zap/ - before/after reports, scan configuration, and remediation evidence.",
            "monitoring/ - Compose, Prometheus, Grafana provisioning, dashboard, and screenshot.",
            "deliverables/ - this report in PDF and Word plus the Brightspace checklist.",
        ],
        compact=True,
    )


def add_conclusion(doc):
    add_section_title(doc, "9", "Conclusion", "Submission readiness")
    doc.add_paragraph(
        "StudyFlow now has a smaller initial client bundle, fewer unnecessary React renders, a "
        "faster indexed and cached read path, stronger browser security headers, and continuous "
        "operational monitoring. The same JMeter workload demonstrates meaningful latency gains "
        "without introducing errors."
    )
    add_callout(
        doc,
        "Assignment status: complete",
        "Every rubric area has implementation evidence, measured results, and a reproducible path "
        "in the repository. The final report is ready for Brightspace after the completed repository "
        "is pushed to the GitHub URL below.",
    )
    doc.add_heading("Repository", level=2)
    p = doc.add_paragraph()
    add_hyperlink(
        p,
        "GitHub repository - StudyFlow Assignment 3",
        "https://github.com/beaprogram/studyflow-assignment3",
    )
    doc.add_heading("Primary Brightspace deliverable", level=2)
    p = doc.add_paragraph()
    r = p.add_run("StudyFlow_Assignment3_Report.pdf")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph(
        "The Word file is retained as an editable backup. Source artifacts are submitted through "
        "the repository link in this report; secrets, node_modules, and generated build folders are excluded."
    )
    doc.add_heading("Evidence index", level=2)
    add_table(
        doc,
        ["Question", "Go to"],
        [
            ["How was load tested?", "Section 2 and jmeter/"],
            ["What changed on the client?", "Section 3 and client/"],
            ["What changed on the server?", "Section 4 and api/"],
            ["What improved?", "Section 5 and PERFORMANCE_EVIDENCE.md"],
            ["Were security findings fixed?", "Section 6 and zap/"],
            ["How is the system monitored?", "Section 7 and monitoring/"],
        ],
        [2.90, 3.60],
        font_size=8.8,
    )


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_executive_summary(doc)
    add_methodology(doc)
    add_client_optimizations(doc)
    add_server_optimizations(doc)
    add_results(doc)
    add_security(doc)
    add_monitoring(doc)
    add_validation(doc)
    add_conclusion(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
